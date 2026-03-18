import re
import pdfplumber
import fitz  # pymupdf
from docx import Document


def normalize_url(url: str) -> str:
    url = url.replace("http://www.",        "https://www.")
    url = url.replace("http://linkedin.com",  "https://www.linkedin.com")
    url = url.replace("https://linkedin.com", "https://www.linkedin.com")
    url = url.replace("http://github.com",    "https://github.com")
    return url.rstrip("/")


def _classify_url(url: str) -> str:
    url_lower = url.lower()
    if "github.com"   in url_lower: return "github"
    if "linkedin.com" in url_lower: return "linkedin"
    return "other"


def _extract_urls_from_text(text: str) -> list:
    return re.findall(r'https?://[^\s\'"<>]+', text)


def _extract_email_from_text(text: str) -> str:
    match = re.search(r'[\w.+-]+@[\w-]+\.[a-zA-Z]{2,}', text)
    return match.group(0) if match else ""


def _extract_from_docx(docx_path: str) -> dict:
    doc = Document(docx_path)
 
    # extract paragraphs
    lines = [para.text for para in doc.paragraphs if para.text.strip()]
 
    # extract table cells — many resumes use table layouts
    for table in doc.tables:
        seen_cells = set()
        for row in table.rows:
            for cell in row.cells:
                # deduplicate merged cells
                cell_id = id(cell._tc)
                if cell_id in seen_cells:
                    continue
                seen_cells.add(cell_id)
                cell_text = cell.text.strip()
                if cell_text:
                    lines.append(cell_text)
 
    text = "\n".join(lines)
 
    links = []
    # read hyperlinks from XML relationship layer
    for rel in doc.part.rels.values():
        if "hyperlink" in rel.reltype:
            links.append(rel.target_ref)
    # regex fallback for plain text URLs
    links += _extract_urls_from_text(text)
 
    seen = set()
    github_url = linkedin_url = ""
    for raw in links:
        url  = normalize_url(raw)
        if url in seen: continue
        seen.add(url)
        kind = _classify_url(url)
        if kind == "github"   and not github_url:   github_url   = url
        if kind == "linkedin" and not linkedin_url: linkedin_url = url
 
    return {
        "text":         text,
        "github_url":   github_url,
        "linkedin_url": linkedin_url,
        "email":        _extract_email_from_text(text),
    }
 
 
def _extract_from_pdf(pdf_path: str) -> dict:
    text  = ""
    links = []

    # pdfplumber — text + annotation URIs
    try:
        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                text += page.extract_text() or ""
                for annot in (page.annots or []):
                    uri = annot.get("uri", "")
                    if uri: links.append(uri)
    except Exception:
        pass

    # pymupdf — annotation layer (more reliable for hyperlinks)
    try:
        doc = fitz.open(pdf_path)
        for page in doc:
            for link in page.get_links():
                uri = link.get("uri", "")
                if uri: links.append(uri)
        doc.close()
    except Exception:
        pass

    # regex fallback
    links += _extract_urls_from_text(text)

    seen = set()
    github_url = linkedin_url = ""
    for raw in links:
        url  = normalize_url(raw)
        if url in seen: continue
        seen.add(url)
        kind = _classify_url(url)
        if kind == "github"   and not github_url:   github_url   = url
        if kind == "linkedin" and not linkedin_url: linkedin_url = url

    return {
        "text":         text,
        "github_url":   github_url,
        "linkedin_url": linkedin_url,
        "email":        _extract_email_from_text(text),
    }


def parse_resume(path: str) -> dict:
    """Main entry point — dispatches to docx or pdf parser."""
    if path.endswith(".docx"):
        return _extract_from_docx(path)
    return _extract_from_pdf(path)